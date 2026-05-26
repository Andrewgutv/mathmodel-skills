from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document


DEFAULT_OUTPUT = "论文_模板一致版.docx"


@dataclass
class WordEnvironmentStatus:
    ok: bool
    category: str
    message: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Merge a full body DOCX into a teacher template while preserving front matter."
    )
    parser.add_argument("source_docx", help="Path to 文档/论文_完整终稿.docx")
    parser.add_argument(
        "--template",
        required=True,
        help="Path to the teacher template DOCX or DOC file.",
    )
    parser.add_argument(
        "--output",
        default=DEFAULT_OUTPUT,
        help="Output DOCX filename, written next to the source DOCX by default.",
    )
    return parser.parse_args()


def get_title(source_docx: Path) -> str:
    document = Document(source_docx)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            return text
    raise ValueError(f"No non-empty paragraph found in source DOCX: {source_docx}")


def ps_quote(path: Path) -> str:
    return str(path).replace("'", "''")


def detect_word_environment() -> WordEnvironmentStatus:
    if sys.platform != "win32":
        return WordEnvironmentStatus(False, "UNSUPPORTED_PLATFORM", "Template merge currently requires Windows.")

    probe = subprocess.run(
        [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-Command",
            "$word = $null; try { $word = New-Object -ComObject Word.Application; 'WORD_COM_OK' } catch { 'WORD_COM_FAIL:' + $_.Exception.Message } finally { if ($word -ne $null) { $word.Quit() } }",
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )
    stdout = probe.stdout.strip()
    if "WORD_COM_OK" in stdout:
        return WordEnvironmentStatus(True, "OK", "Word COM automation is available.")
    if stdout.startswith("WORD_COM_FAIL:"):
        return WordEnvironmentStatus(False, "WORD_COM_UNAVAILABLE", stdout.removeprefix("WORD_COM_FAIL:").strip())
    return WordEnvironmentStatus(False, "WORD_COM_UNKNOWN", stdout or probe.stderr.strip() or "Unknown Word COM probe failure.")


def inspect_template_compatibility(template_path: Path) -> tuple[bool, list[str]]:
    notes: list[str] = []
    try:
        document = Document(template_path)
    except Exception as exc:
        return False, [f"Template is unreadable as DOCX: {exc}"]

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    styles = [p.style.name for p in document.paragraphs]

    has_abstract = any(text == "摘要" for text in paragraphs)
    has_title_placeholder = any(text == "论文题目" for text in paragraphs)
    has_toc = any("目录" in text for text in paragraphs) or "TOC Heading" in styles or any(style.startswith("toc ") for style in styles)

    if not has_abstract:
        notes.append("Template does not contain a detectable 摘要 heading.")
    if not has_title_placeholder:
        notes.append("Template does not contain a detectable 论文题目 placeholder.")
    if not has_toc:
        notes.append("Template does not contain detectable 目录/TOC front matter.")

    return has_abstract, notes


def build_powershell_script(source_docx: Path, template_path: Path, output_docx: Path, source_title: str) -> str:
    src = ps_quote(source_docx)
    template = ps_quote(template_path)
    out = ps_quote(output_docx)
    title = source_title.replace("'", "''")
    return f"""
$ErrorActionPreference = 'Stop'
$src = '{src}'
$template = '{template}'
$out = '{out}'
$sourceTitle = '{title}'
Copy-Item -LiteralPath $template -Destination $out -Force
$word = $null
$doc = $null
$srcDoc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open($out)
  $srcDoc = $word.Documents.Open($src, $false, $true)

  $srcAbstractStart = $null
  for ($i = 1; $i -le $srcDoc.Paragraphs.Count; $i++) {{
    $p = $srcDoc.Paragraphs.Item($i).Range
    $txt = $p.Text.Replace("`r",' ').Replace("`a",' ').Trim()
    if ($txt -eq '摘要') {{
      try {{
        $styleName = $p.Style.NameLocal
      }} catch {{
        $styleName = ''
      }}
      if ($styleName -eq '标题 1' -or $styleName -eq 'Heading 1') {{
        $srcAbstractStart = $p.Start
        break
      }}
    }}
  }}
  if ($null -eq $srcAbstractStart) {{ throw 'Source abstract heading not found.' }}

  $tplAbstractStart = $null
  for ($i = 1; $i -le $doc.Paragraphs.Count; $i++) {{
    $p = $doc.Paragraphs.Item($i).Range
    $txt = $p.Text.Replace("`r",' ').Replace("`a",' ').Trim()
    if ($txt -eq '摘要') {{
      try {{
        $styleName = $p.Style.NameLocal
      }} catch {{
        $styleName = ''
      }}
      if ($styleName -eq '标题 1' -or $styleName -eq 'Heading 1') {{
        $tplAbstractStart = $p.Start
        break
      }}
    }}
  }}
  if ($null -eq $tplAbstractStart) {{ throw 'Template abstract heading not found.' }}

  $titleReplaced = 0
  for ($i = 1; $i -le $doc.Paragraphs.Count; $i++) {{
    $p = $doc.Paragraphs.Item($i).Range
    $txt = $p.Text.Replace("`r",' ').Replace("`a",' ').Trim()
    if ($txt -eq '论文题目') {{
      $p.Text = $sourceTitle + "`r"
      $titleReplaced++
      if ($titleReplaced -ge 2) {{ break }}
    }}
  }}

  $deleteRange = $doc.Range($tplAbstractStart, $doc.Content.End - 1)
  $deleteRange.Delete()

  $insertRange = $doc.Range($doc.Content.End - 1, $doc.Content.End - 1)
  $copyRange = $srcDoc.Range($srcAbstractStart, $srcDoc.Content.End - 1)
  $insertRange.FormattedText = $copyRange.FormattedText

  foreach ($toc in $doc.TablesOfContents) {{
    $toc.Update()
  }}
  $doc.Fields.Update() | Out-Null
  $doc.Repaginate()
  $doc.Save()
  $pages = $doc.ComputeStatistics(2)
  Write-Output ("OK|PAGES=" + $pages)
}}
finally {{
  if ($srcDoc -ne $null) {{
    $srcDoc.Close([ref]0)
    [void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($srcDoc)
  }}
  if ($doc -ne $null) {{
    $doc.Close([ref]0)
    [void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($doc)
  }}
  if ($word -ne $null) {{
    $word.Quit()
    [void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($word)
  }}
  [GC]::Collect()
  [GC]::WaitForPendingFinalizers()
}}
"""


def run_merge(source_docx: Path, template_path: Path, output_docx: Path, source_title: str) -> subprocess.CompletedProcess[str]:
    script = build_powershell_script(source_docx, template_path, output_docx, source_title)
    return subprocess.run(
        [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-Command",
            script,
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )


def inspect_output(output_docx: Path) -> tuple[list[str], list[str]]:
    critical_errors: list[str] = []
    review_warnings: list[str] = []
    notes: list[str] = []

    if not output_docx.exists() or output_docx.stat().st_size == 0:
        return [f"Output DOCX missing or empty: {output_docx}"], notes

    try:
        document = Document(output_docx)
    except Exception as exc:
        return [f"Template-aligned DOCX is unreadable: {output_docx} ({exc})"], notes

    abstract_count = sum(
        1
        for p in document.paragraphs
        if p.text.strip() == "摘要" and p.style.name.startswith("Heading")
    )
    if abstract_count != 1:
        review_warnings.append(f"Expected exactly one '摘要' heading, found {abstract_count}.")

    first_texts = [p.text.strip() for p in document.paragraphs[:120] if p.text.strip()]
    first_styles = [p.style.name for p in document.paragraphs[:120]]
    if not any("承诺书" in text for text in first_texts):
        review_warnings.append("Template front matter check failed: no 承诺书 text detected in the leading paragraphs.")
    if not (
        any("目录" in text for text in first_texts)
        or "TOC Heading" in first_styles
        or any(style.startswith("toc ") for style in first_styles)
    ):
        review_warnings.append("Template front matter check failed: no 目录 text detected in the leading paragraphs.")

    notes.append(f"Paragraphs: {len(document.paragraphs)}")
    notes.append(f"Tables: {len(document.tables)}")
    if review_warnings:
        notes.extend(f"REVIEW: {warning}" for warning in review_warnings)
    return critical_errors, notes


def determine_merge_status(errors: list[str], notes: list[str]) -> str:
    if errors:
        return "FAIL"
    if any(note.startswith("REVIEW:") for note in notes):
        return "REVIEW_REQUIRED"
    return "PASS"


def main() -> int:
    args = parse_args()
    source_docx = Path(args.source_docx).expanduser().resolve()
    template_path = Path(args.template).expanduser().resolve()
    output_docx = (source_docx.parent / args.output).resolve()

    if not source_docx.exists():
        raise SystemExit(f"Source DOCX not found: {source_docx}")
    if not template_path.exists():
        raise SystemExit(f"Template file not found: {template_path}")
    if output_docx == source_docx:
        raise SystemExit("Output path must be different from source DOCX.")

    env_status = detect_word_environment()
    if not env_status.ok:
        print("MERGE_STATUS: FAIL")
        print(f"FAIL_CATEGORY: {env_status.category}")
        print(f"DETAIL: {env_status.message}")
        return 1

    template_ok, template_notes = inspect_template_compatibility(template_path)
    if not template_ok:
        print("MERGE_STATUS: FAIL")
        print("FAIL_CATEGORY: TEMPLATE_INCOMPATIBLE")
        for note in template_notes:
            print(f"DETAIL: {note}")
        return 1

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    source_title = get_title(source_docx)

    completed = run_merge(source_docx, template_path, output_docx, source_title)
    if completed.stdout.strip():
        print(completed.stdout.strip())
    if completed.returncode != 0:
        print("MERGE_STATUS: FAIL")
        print("FAIL_CATEGORY: WORD_MERGE_RUNTIME")
        if completed.stderr.strip():
            print(completed.stderr.strip(), file=sys.stderr)
        return completed.returncode

    errors, notes = inspect_output(output_docx)
    status = determine_merge_status(errors, notes)
    print(f"MERGE_STATUS: {status}")
    if errors:
        print("FAIL_CATEGORY: OUTPUT_VALIDATION")
        for error in errors:
            print(error, file=sys.stderr)
        raise SystemExit(1)

    print(f"Generated template-aligned DOCX: {output_docx}")
    for note in template_notes + notes:
        print(f"- {note}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
