from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_OUTPUT = "论文_完整终稿.docx"
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".gif"}


@dataclass
class Block:
    kind: str
    text: str
    meta: dict[str, str] = field(default_factory=dict)


@dataclass
class RenderState:
    equation_index: int = 0
    figure_index: int = 0
    table_index: int = 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a stable full-body DOCX from 文档/论文.md."
    )
    parser.add_argument("markdown", help="Path to 文档/论文.md")
    parser.add_argument(
        "--output",
        default=DEFAULT_OUTPUT,
        help="Output DOCX filename, written next to the markdown file by default.",
    )
    parser.add_argument(
        "--template",
        default=None,
        help="Reserved for future use. The modern full-body builder currently ignores this option.",
    )
    parser.add_argument(
        "--no-template-search",
        action="store_true",
        help="Accepted for compatibility. The modern full-body builder does not use a template.",
    )
    parser.add_argument(
        "--title",
        default="数学建模竞赛论文",
        help="Reserved for future use.",
    )
    return parser.parse_args()


def strip_internal_blocks(markdown_text: str) -> str:
    patterns = [
        r"(?s)<!-- AUTO_RESULT_SUMMARY -->.*?<!-- AUTO_QUESTION_RESULTS_END -->",
        r"(?s)<!-- AUTO_DEFENSE_RESULTS_START -->.*?<!-- AUTO_DEFENSE_RESULTS_END -->",
        r"(?m)^<!-- .*? -->\s*$",
    ]
    cleaned = markdown_text
    for pattern in patterns:
        cleaned = re.sub(pattern, "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip() + "\n"


def split_blocks(markdown_text: str) -> list[Block]:
    blocks: list[Block] = []
    current: list[str] = []
    in_code = False

    for line in markdown_text.splitlines():
        if line.strip().startswith("```"):
            if current:
                blocks.append(Block("paragraph", "\n".join(current).strip()))
                current = []
            in_code = not in_code
            if in_code:
                current = [line]
            else:
                current.append(line)
                blocks.append(Block("code", "\n".join(current).strip()))
                current = []
            continue

        if in_code:
            current.append(line)
            continue

        if not line.strip():
            if current:
                blocks.append(classify_block("\n".join(current).strip()))
                current = []
            continue

        current.append(line)

    if current:
        blocks.append(classify_block("\n".join(current).strip()))

    return [block for block in blocks if block.text]


def classify_block(text: str) -> Block:
    stripped = text.strip()
    if stripped.startswith("# "):
        return Block("heading1", stripped[2:].strip())
    if stripped.startswith("## "):
        return Block("heading2", stripped[3:].strip())
    if stripped.startswith("### "):
        return Block("heading3", stripped[4:].strip())
    if is_display_math(stripped):
        return Block("equation", stripped)
    if is_pipe_table(stripped):
        return Block("table", stripped)
    if all(is_list_item(line) for line in stripped.splitlines()):
        return Block("list", stripped)
    if looks_like_image_line(stripped):
        return Block("image", stripped, extract_image_meta(stripped))
    return Block("paragraph", stripped)


def is_pipe_table(text: str) -> bool:
    lines = text.splitlines()
    if len(lines) < 2:
        return False
    if "|" not in lines[0]:
        return False
    return bool(re.match(r"^\s*\|?[\-\:\s|]+\|?\s*$", lines[1]))


def is_list_item(line: str) -> bool:
    return bool(re.match(r"^\s*(?:[-*]|\d+\.)\s+", line))


def is_display_math(text: str) -> bool:
    return text.startswith("$$") and text.endswith("$$") and len(text) >= 4


def looks_like_image_line(text: str) -> bool:
    if re.match(r"^!\[.*?\]\(.*?\)$", text):
        return True
    lower = text.lower()
    return any(lower.endswith(ext) for ext in IMAGE_EXTENSIONS)


def extract_image_meta(text: str) -> dict[str, str]:
    match = re.match(r"^!\[(.*?)\]\((.*?)\)$", text.strip())
    if not match:
        return {}
    return {"caption": match.group(1).strip(), "path": match.group(2).strip()}


def create_document() -> Document:
    document = Document()
    configure_document(document)
    apply_global_styles(document)
    return document


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)


def apply_global_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except AttributeError:
        pass
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, bold, center in (
        ("Heading 1", 15, True, True),
        ("Heading 2", 14, True, False),
        ("Heading 3", 12, True, False),
    ):
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        try:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style_name == "Heading 1" else "宋体")
        except AttributeError:
            pass
        if center:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_run_font(run, east_asia: str = "宋体", western: str = "Times New Roman", size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = western
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    except AttributeError:
        pass


def render_blocks(document: Document, blocks: list[Block], project_root: Path) -> None:
    state = RenderState()
    for block in blocks:
        if block.kind.startswith("heading"):
            level = int(block.kind[-1])
            paragraph = document.add_paragraph(style=f"Heading {level}")
            run = paragraph.add_run(block.text)
            set_run_font(run, east_asia="黑体" if level == 1 else "宋体", size=15 if level == 1 else 14 if level == 2 else 12, bold=True)
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if block.kind == "paragraph":
            add_paragraph(document, block.text)
            continue

        if block.kind == "list":
            add_list(document, block.text)
            continue

        if block.kind == "equation":
            add_equation(document, block.text, state)
            continue

        if block.kind == "table":
            add_table(document, block.text, state)
            continue

        if block.kind == "image":
            add_image(document, block, project_root, state)
            continue

        if block.kind == "code":
            add_code_block(document, block.text)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for segment, is_math in split_inline_math(text):
        run = paragraph.add_run(segment)
        set_run_font(run, italic=is_math)


def split_inline_math(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    cursor = 0
    for match in re.finditer(r"\$(.+?)\$", text):
        if match.start() > cursor:
            parts.append((text[cursor:match.start()], False))
        parts.append((match.group(1), True))
        cursor = match.end()
    if cursor < len(text):
        parts.append((text[cursor:], False))
    return parts or [(text, False)]


def add_list(document: Document, text: str) -> None:
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.37)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = paragraph.add_run(stripped)
        set_run_font(run)


def parse_pipe_table(text: str) -> list[list[str]]:
    rows = []
    for index, line in enumerate(text.splitlines()):
        if index == 1 and re.match(r"^\s*\|?[\-\:\s|]+\|?\s*$", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(document: Document, text: str, state: RenderState) -> None:
    rows = parse_pipe_table(text)
    if not rows:
        return
    state.table_index += 1
    caption = f"表{state.table_index}"
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=10, bold=True)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, cell_text in enumerate(row):
            table.cell(row_index, col_index).text = cell_text
            for paragraph in table.cell(row_index, col_index).paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=10, bold=row_index == 0)


def add_image(document: Document, block: Block, project_root: Path, state: RenderState) -> None:
    image_ref = block.meta.get("path")
    if not image_ref:
        add_paragraph(document, block.text)
        return

    image_path = (project_root / image_ref).resolve()
    if not image_path.exists():
        image_path = Path(image_ref).expanduser()
    if not image_path.exists():
        add_paragraph(document, f"[Missing image] {image_ref}")
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(14.2))
    state.figure_index += 1
    caption = block.meta.get("caption", "").strip()
    caption_text = f"图{state.figure_index}" if not caption else f"图{state.figure_index} {caption}"
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption_text)
    set_run_font(caption_run, size=10)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)


def add_equation(document: Document, text: str, state: RenderState) -> None:
    state.equation_index += 1
    equation_body = text.strip()[2:-2].strip()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)

    equation_run = paragraph.add_run(equation_body)
    set_run_font(equation_run, east_asia="Cambria Math", western="Cambria Math", size=12, italic=True)

    number_run = paragraph.add_run(f"    ({state.equation_index})")
    set_run_font(number_run, east_asia="Cambria Math", western="Cambria Math", size=12)


def save_document(document: Document, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    args = parse_args()
    markdown_path = Path(args.markdown).expanduser().resolve()
    if not markdown_path.exists():
        raise SystemExit(f"Markdown file not found: {markdown_path}")

    markdown_text = markdown_path.read_text(encoding="utf-8")
    markdown_text = strip_internal_blocks(markdown_text)
    blocks = split_blocks(markdown_text)

    document = create_document()
    project_root = markdown_path.parent.parent
    render_blocks(document, blocks, project_root)

    output_path = markdown_path.parent / args.output
    save_document(document, output_path)
    if not output_path.exists():
        raise SystemExit(f"DOCX was not written as expected: {output_path}")

    print(f"Generated DOCX: {output_path}")
    print("Template used: none")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
