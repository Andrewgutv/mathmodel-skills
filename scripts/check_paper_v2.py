from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path
from zipfile import BadZipFile

from docx import Document


REQUIRED_HEADINGS = [
    "# 摘要",
    "# 一、问题的提出",
    "# 二、问题的分析",
    "# 三、模型假设",
    "# 四、符号说明",
    "# 五、建模与求解",
    "# 六、模型的检验/敏感性分析",
    "# 七、模型的评价与推广",
    "# 参考文献",
]
AI_DECLARATION_PATTERNS = [
    "AI使用声明",
    "AI 使用声明",
    "人工智能使用声明",
]
PLACEHOLDER_TOKENS = ("待填写", "示例指标")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Run structural and project-level checks for a mathmodel project."
    )
    parser.add_argument("markdown", help="Path to 文档/论文.md")
    parser.add_argument(
        "--project-root",
        default=None,
        help="Optional project root used for checking code, results, figures, and delivery files.",
    )
    parser.add_argument(
        "--mode",
        choices=("paper", "full_delivery", "defense"),
        default="paper",
        help="Select which deliverables are mandatory for this check run.",
    )
    parser.add_argument(
        "--require-ai-declaration",
        action="store_true",
        help="When set, fail if the paper does not include an AI usage declaration.",
    )
    return parser.parse_args()


def load_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists() or path.stat().st_size == 0:
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def csv_looks_like_placeholder(path: Path) -> bool:
    try:
        content = path.read_text(encoding="utf-8-sig")
    except OSError:
        return True
    return any(token in content for token in PLACEHOLDER_TOKENS)


def find_missing_headings(text: str) -> list[str]:
    return [heading for heading in REQUIRED_HEADINGS if heading not in text]


def has_ai_declaration(text: str) -> bool:
    return any(pattern in text for pattern in AI_DECLARATION_PATTERNS)


def extract_section(text: str, heading: str) -> str:
    pattern = rf"{re.escape(heading)}\n(.*?)(?:\n# |\Z)"
    match = re.search(pattern, text, flags=re.DOTALL)
    return match.group(1).strip() if match else ""


def has_reference_items(text: str) -> bool:
    reference_section = extract_section(text, "# 参考文献")
    if not reference_section:
        return False
    return bool(re.search(r"^\s*(?:[-*]|\d+\.)\s+", reference_section, re.MULTILINE))


def expected_question_numbers(project_root: Path) -> list[int]:
    mapping_path = project_root / "文档" / "问题-方法对照表.md"
    if not mapping_path.exists():
        return []
    text = mapping_path.read_text(encoding="utf-8")
    numbers = set()
    for match in re.findall(r"##\s+问题(\d+)", text):
        numbers.add(int(match))
    for match in re.findall(r"^\|\s*问题(\d+)\s*\|", text, flags=re.MULTILINE):
        numbers.add(int(match))
    return sorted(numbers)


def count_markdown_images(text: str) -> int:
    return len(re.findall(r"!\[.*?\]\(.*?\)", text))


def count_figure_refs(text: str) -> int:
    return len(re.findall(r"图\d+", text))


def count_table_refs(text: str) -> int:
    return len(re.findall(r"表\d+", text))


def count_formula_markers(text: str) -> int:
    display_math = len(re.findall(r"\$\$.*?\$\$", text, flags=re.DOTALL))
    inline_math = len(re.findall(r"\$[^$\n]+\$", text))
    return display_math + inline_math


def table_inventory_has_placeholders(path: Path) -> bool:
    if not path.exists():
        return False
    text = path.read_text(encoding="utf-8", errors="replace")
    return "待补充" in text


def list_real_figure_files(image_dir: Path) -> list[Path]:
    return [
        path
        for path in image_dir.iterdir()
        if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp", ".gif"}
    ] if image_dir.exists() else []


def has_result_evidence(markdown_text: str) -> bool:
    patterns = [
        r"结果/",
        r"综合结果汇总\.csv",
        r"问题\d+_结果汇总\.csv",
        r"自动回填结果摘要",
        r"对于问题\d+",
    ]
    return any(re.search(pattern, markdown_text) for pattern in patterns)


def inspect_docx(path: Path, template_expected: bool) -> tuple[list[str], list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    notes: list[str] = []

    if path.stat().st_size == 0:
        return [f"DOCX is empty: {path}"], warnings, notes

    try:
        document = Document(path)
    except BadZipFile:
        return [f"DOCX is unreadable or corrupted: {path}"], warnings, notes
    except Exception as exc:
        return [f"DOCX inspection failed for {path}: {exc}"], warnings, notes

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    notes.append(f"DOCX checked: {path.name} (paragraphs={len(document.paragraphs)}, tables={len(document.tables)})")

    abstract_count = sum(
        1
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "摘要" and paragraph.style.name.startswith("Heading")
    )
    if abstract_count != 1:
        if template_expected:
            warnings.append(f"{path.name}: expected exactly one '摘要', found {abstract_count}.")
        else:
            errors.append(f"{path.name}: expected exactly one '摘要', found {abstract_count}.")

    if template_expected:
        leading = paragraphs[:120]
        if not any("承诺书" in text for text in leading):
            warnings.append(f"{path.name}: no 承诺书 text detected in front matter.")
        if not any("目录" in text for text in leading):
            warnings.append(f"{path.name}: no 目录 text detected in front matter.")

    return errors, warnings, notes


def check_model_layer(project_root: Path) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    notes: list[str] = []
    code_dir = project_root / "代码"

    main_script = code_dir / "main.py"
    run_all_script = code_dir / "run_all.py"
    question_scripts = sorted(code_dir.glob("code_问题*_*.py"))
    if not (main_script.exists() or run_all_script.exists() or question_scripts):
        errors.append("Model layer: no solver entry found under 代码/.")
    else:
        notes.append("Model layer: solver entry detected.")

    result_dir = project_root / "结果"
    expected_numbers = expected_question_numbers(project_root)
    if expected_numbers:
        notes.append(f"Model layer: expected question count = {len(expected_numbers)}.")
    else:
        notes.append("Model layer: expected question count could not be inferred from 问题-方法对照表.md.")

    question_results = sorted(result_dir.glob("问题*_结果汇总.csv"))
    if not question_results:
        errors.append("Model layer: no per-question result CSVs found under 结果/.")
    else:
        notes.append(f"Model layer: detected {len(question_results)} per-question result CSV file(s).")
        for path in question_results:
            rows = load_csv_rows(path)
            if not rows:
                errors.append(f"Model layer: result CSV is empty or header-only: {path}")
            elif csv_looks_like_placeholder(path):
                errors.append(f"Model layer: result CSV still looks like a placeholder: {path}")

    if expected_numbers:
        missing = [number for number in expected_numbers if not (result_dir / f"问题{number}_结果汇总.csv").exists()]
        if missing:
            formatted = ", ".join(str(number) for number in missing)
            errors.append(f"Model layer: missing expected per-question result CSVs for: {formatted}")

    combined = result_dir / "综合结果汇总.csv"
    combined_rows = load_csv_rows(combined)
    if not combined.exists():
        errors.append("Model layer: 综合结果汇总.csv is missing.")
    elif not combined_rows:
        errors.append("Model layer: 综合结果汇总.csv is empty or header-only.")
    elif csv_looks_like_placeholder(combined):
        errors.append("Model layer: 综合结果汇总.csv still looks like a placeholder.")
    else:
        notes.append(f"Model layer: 综合结果汇总.csv has {len(combined_rows)} row(s).")

    return errors, notes


def check_tables_figures_formulas(project_root: Path, markdown_text: str) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    notes: list[str] = []

    figure_inventory = project_root / "图片" / "图表清单.md"
    if not figure_inventory.exists():
        errors.append("Object layer: 图片/图表清单.md is missing.")
    else:
        notes.append("Object layer: figure inventory exists.")
        inventory_text = figure_inventory.read_text(encoding="utf-8", errors="replace")
        if "待补充" in inventory_text and count_figure_refs(markdown_text) > 0:
            errors.append("Object layer: figure inventory still contains placeholders while the paper references figures.")

    figure_refs = count_figure_refs(markdown_text)
    table_refs = count_table_refs(markdown_text)
    formula_markers = count_formula_markers(markdown_text)
    markdown_images = count_markdown_images(markdown_text)
    real_figure_files = list_real_figure_files(project_root / "图片")
    notes.append(f"Object layer: figure refs = {figure_refs}")
    notes.append(f"Object layer: table refs = {table_refs}")
    notes.append(f"Object layer: formula markers = {formula_markers}")
    notes.append(f"Object layer: markdown images = {markdown_images}")
    notes.append(f"Object layer: real figure files = {len(real_figure_files)}")

    if markdown_images > 0 and figure_refs == 0:
        errors.append("Object layer: markdown images exist, but no 图N references were found.")

    if figure_refs > 0 and not real_figure_files:
        errors.append("Object layer: the paper references figures, but no real image files were found under 图片/.")

    if figure_refs > 0 and table_inventory_has_placeholders(figure_inventory):
        errors.append("Object layer: the paper references figures, but 图表清单.md still contains placeholders.")

    if table_refs > 0 and table_inventory_has_placeholders(figure_inventory):
        errors.append("Object layer: the paper references tables, but 图表清单.md still contains placeholders.")

    if table_refs > 0 and "|" not in markdown_text:
        errors.append("Object layer: the paper references tables, but no markdown table structure was detected.")

    if formula_markers == 0 and "公式" in markdown_text:
        errors.append("Object layer: the paper mentions formulas but no formula markers were detected.")

    return errors, notes


def check_document_layer(project_root: Path, markdown_text: str, mode: str) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    notes: list[str] = []

    if "<!-- AUTO_RESULT_SUMMARY -->" not in markdown_text:
        errors.append("Document layer: 论文.md is missing the AUTO_RESULT_SUMMARY marker.")
    else:
        notes.append("Document layer: result summary marker detected in 论文.md.")

    if not has_result_evidence(markdown_text):
        errors.append("Document layer: 论文.md does not appear to contain result-driven content.")

    if mode in {"full_delivery", "defense"}:
        defense_prep = project_root / "答辩" / "答辩准备.md"
        if not defense_prep.exists():
            errors.append("Document layer: 答辩/答辩准备.md is missing.")
        else:
            notes.append("Document layer: defense prep exists.")

    return errors, notes


def check_delivery_layer(project_root: Path, mode: str) -> tuple[list[str], list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    notes: list[str] = []

    expected_paths = [
        project_root / "文档" / "方案.md",
        project_root / "文档" / "解题思路.md",
        project_root / "文档" / "解题过程.md",
        project_root / "文档" / "论文.md",
    ]
    if mode in {"full_delivery", "defense"}:
        expected_paths.extend(
            [
                project_root / "答辩" / "答辩准备.md",
                project_root / "答辩" / "答辩讲稿.md",
                project_root / "答辩" / "PPT设计文案.md",
                project_root / "答辩" / "老师高频问题与回答.md",
            ]
        )

    missing = [path for path in expected_paths if not path.exists()]
    if missing:
        errors.append("Delivery layer: missing expected project deliverables:")
        errors.extend(f"- {path}" for path in missing)
    else:
        notes.append("Delivery layer: required deliverables for current mode are present.")

    full_docx = project_root / "文档" / "论文_完整终稿.docx"
    if full_docx.exists():
        docx_errors, docx_warnings, docx_notes = inspect_docx(full_docx, template_expected=False)
        errors.extend(docx_errors)
        warnings.extend(docx_warnings)
        notes.extend(docx_notes)
    elif mode == "full_delivery":
        errors.append("Delivery layer: 论文_完整终稿.docx is missing in full_delivery mode.")

    template_docx = project_root / "文档" / "论文_模板一致版.docx"
    if template_docx.exists():
        notes.append(f"Delivery layer: template-aligned DOCX detected: {template_docx}")
        docx_errors, docx_warnings, docx_notes = inspect_docx(template_docx, template_expected=True)
        errors.extend(docx_errors)
        warnings.extend(docx_warnings)
        notes.extend(docx_notes)
    else:
        notes.append("Delivery layer: template-aligned DOCX not found; acceptable when no teacher template was requested.")

    return errors, warnings, notes


def determine_status(errors: list[str], warnings: list[str], notes: list[str]) -> str:
    if not errors:
        if warnings or any("template-aligned DOCX not found" in note for note in notes):
            return "PASS_WITH_TEMPLATE_WARNING"
        return "PASS"

    if any(error.startswith("Model layer:") for error in errors):
        return "FAIL_MODEL"
    if any(error.startswith("Object layer:") or error.startswith("Document layer:") or error.startswith("Paper structure:") for error in errors):
        return "FAIL_DOCUMENT"
    if any(error.startswith("Delivery layer:") or "DOCX" in error or "论文_完整终稿.docx" in error for error in errors):
        return "FAIL_WORD"
    return "FAIL"


def main() -> int:
    args = parse_args()
    markdown_path = Path(args.markdown).resolve()
    text = markdown_path.read_text(encoding="utf-8")

    errors: list[str] = []
    warnings: list[str] = []
    notes: list[str] = []

    missing_headings = find_missing_headings(text)
    if missing_headings:
        errors.append("Paper structure: missing required headings:")
        errors.extend(f"- {heading}" for heading in missing_headings)

    if args.require_ai_declaration and not has_ai_declaration(text):
        errors.append("Paper structure: missing required AI usage declaration near the end of the paper.")

    if not has_reference_items(text):
        errors.append("Paper structure: 参考文献 section exists but no list-style reference items were detected.")

    if args.project_root:
        project_root = Path(args.project_root).resolve()
        layer_errors, layer_notes = check_model_layer(project_root)
        errors.extend(layer_errors)
        notes.extend(layer_notes)
        layer_errors, layer_notes = check_document_layer(project_root, text, args.mode)
        errors.extend(layer_errors)
        notes.extend(layer_notes)
        layer_errors, layer_notes = check_tables_figures_formulas(project_root, text)
        errors.extend(layer_errors)
        notes.extend(layer_notes)
        layer_errors, layer_warnings, layer_notes = check_delivery_layer(project_root, args.mode)
        errors.extend(layer_errors)
        warnings.extend(layer_warnings)
        notes.extend(layer_notes)

    status = determine_status(errors, warnings, notes)

    print(f"STATUS: {status}")
    print(f"MODE: {args.mode}")
    if warnings:
        print("Warnings:")
        for item in warnings:
            print(item)
    if errors:
        print("Paper check failed:")
        for item in errors:
            print(item)
        if notes:
            print("\nNotes:")
            for note in notes:
                print("-", note)
        return 1

    print("Paper check passed.")
    for note in notes:
        print("-", note)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
